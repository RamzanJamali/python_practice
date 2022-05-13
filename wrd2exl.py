import docx
import xlrd
import re
import io

from docx import Document
doc = Document('emp.docx')

print(doc.paragraphs[0].text)
#print(doc.paragraphs[0].runs) 
#print(doc.paragraphs[2].text)
#for para in doc.paragraphs: 
#    print(para.text)
